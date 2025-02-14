import streamlit as st
import requests
from fpdf import FPDF
from docx import Document
import openai
import PyPDF2
from docx.shared import Inches
import io
from pptx import Presentation
from pptx.util import Inches, Pt
from pptx.enum.text import PP_ALIGN
import zipfile
import os
from pathlib import Path
import csv


openai.api_key = st.secrets["api"]["OPENAI_API_KEY"]


# Function to fetch response from GPT

def fetch_gpt_response(query):
    try:
        response = openai.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are an expert in the pharmaceutical and medical domain only. Only answer those questions and don't answer any other questions. Dont analyze and summarize pdf if the pdf is not related to medical and pharmaceutical domain."},
                {"role": "user", "content": query},
            ],
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error: {str(e)}"

# Function to extract text from PDF
def extract_text_from_pdf(pdf_file):
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    pdf_text = ""
    for page in pdf_reader.pages:
        pdf_text += page.extract_text()
    return pdf_text


def save_as_scorm_pdf(content, output_folder="scorm_package", scorm_zip_name="scorm_package.zip"):
    # Step 1: Create the SCORM folder structure
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # Save the PDF
    pdf_file_path = os.path.join(output_folder, "content.pdf")
    save_as_pdf(content, pdf_file_path)

    # Step 2: Create the HTML file
    html_file_path = os.path.join(output_folder, "index.html")
    with open(html_file_path, "w", encoding="utf-8") as html_file:
        html_file.write(f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>SCORM Content</title>
        </head>
        <body>
            <h1>Research Content Response</h1>
            <iframe src="content.pdf" width="100%" height="600px"></iframe>
        </body>
        </html>
        """)

    # Step 3: Create the imsmanifest.xml file
    manifest_file_path = os.path.join(output_folder, "imsmanifest.xml")
    with open(manifest_file_path, "w", encoding="utf-8") as manifest_file:
        manifest_file.write(f"""
        <?xml version="1.0" encoding="UTF-8"?>
        <manifest xmlns="http://www.imsglobal.org/xsd/imscp_v1p1"
                  xmlns:adlcp="http://www.adlnet.org/xsd/adlcp_v1p3"
                  xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"
                  xsi:schemaLocation="http://www.imsglobal.org/xsd/imscp_v1p1
                                      http://www.imsglobal.org/xsd/imscp_v1p1.xsd
                                      http://www.adlnet.org/xsd/adlcp_v1p3
                                      http://www.adlnet.org/xsd/adlcp_v1p3.xsd">
            <metadata>
                <schema>ADL SCORM</schema>
                <schemaversion>1.2</schemaversion>
            </metadata>
            <organizations>
                <organization identifier="ORG-1">
                    <title>Research Content</title>
                    <item identifier="ITEM-1" identifierref="RES-1">
                        <title>Research Content Response</title>
                    </item>
                </organization>
            </organizations>
            <resources>
                <resource identifier="RES-1" type="webcontent" href="index.html">
                    <file href="index.html"/>
                    <file href="content.pdf"/>
                </resource>
            </resources>
        </manifest>
        """)

    # Step 4: Zip the SCORM package
    with zipfile.ZipFile(scorm_zip_name, 'w', zipfile.ZIP_DEFLATED) as scorm_zip:
        for foldername, subfolders, filenames in os.walk(output_folder):
            for filename in filenames:
                file_path = os.path.join(foldername, filename)
                arcname = os.path.relpath(file_path, output_folder)
                scorm_zip.write(file_path, arcname)

    # Provide the download button for the SCORM package
    with open(scorm_zip_name, "rb") as scorm_file:
        st.download_button("Download SCORM Package", scorm_file, scorm_zip_name, "application/zip")


def save_as_pdf(content, file_name="response.pdf"):
    pdf = FPDF()
    pdf.add_page()

    # Add the logo
    pdf.image('assets/logo.jpeg', x=10, y=8, w=30)

    # Title of the document
    pdf.set_font("Arial", style='B', size=16)
    pdf.ln(30)
    pdf.cell(200, 10, txt="Research Content Response", ln=True, align='C')
    pdf.ln(10)

    # Add content
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(190, 10, content)

    # Save the PDF
    pdf.output(file_name)


def save_as_scorm_word(content, file_name="scorm_package.zip"):
    # Create an in-memory zip file
    scorm_zip = io.BytesIO()

    with zipfile.ZipFile(scorm_zip, 'w') as zf:
        # Create and add manifest.xml
        manifest_content = """<manifest>
            <metadata>
                <schema>ADL SCORM</schema>
                <schemaversion>1.2</schemaversion>
            </metadata>
            <resources>
                <resource identifier="res1" type="webcontent" href="response.docx">
                    <file href="response.docx"/>
                    <file href="response.html"/>
                </resource>
            </resources>
        </manifest>"""
        zf.writestr("imanifest.xml", manifest_content)

        # Create DOCX file
        docx_buffer = io.BytesIO()
        doc = Document()
        # Add the logo to the Word document
        logo_path = "assets/logo.jpeg"
        if Path(logo_path).is_file():
            doc.add_picture(logo_path, width=Inches(1.5))
        doc.add_paragraph('\n')
        doc.add_paragraph("Research Content Response", style='Heading 1')
        doc.add_paragraph('\n')
        doc.add_paragraph(content)
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        zf.writestr("response.docx", docx_buffer.getvalue())

        # Create HTML file
        html_content = f"""
        <html>
        <head><title>Research Content Response</title></head>
        <body>
        <h1>Research Content Response</h1>
        <p>{content.replace('\n', '<br>')}</p>
        </body>
        </html>
        """
        zf.writestr("index.html", html_content)

    scorm_zip.seek(0)
    return scorm_zip.getvalue()


# Usage in Streamlit
def save_as_scorm_button(content):
    scorm_data = save_as_scorm_word(content)
    st.download_button(
        label="Download SCORM Package",
        data=scorm_data,
        file_name="scorm_package.zip",
        mime="application/zip"
    )


# Set up the page configuration (must be the first command)
st.set_page_config(page_title="PDF Analysis", layout="wide", page_icon="📚")

# Title Section with enhanced visuals
st.markdown(
    """
    <h1 style="text-align: center; font-size: 2.5rem; color: #4A90E2;">📚 AI-Powered PDF Analysis And Content Generation</h1>
    <p style="text-align: center; font-size: 1.1rem; color: #555;">Streamline your content creation process with AI technology. Designed for the <strong>pharmaceutical</strong> and <strong>medical</strong> domains.</p>
    """,
    unsafe_allow_html=True,
)
# Horizontal line
st.markdown("---")

 # PDF Analysis Instructions
with st.expander("1️⃣ **PDF Analysis Instructions**",expanded=True):
    st.markdown("""
    - Analyze and extract content from uploaded PDFs.
    - **Steps**:
    1. Upload a **PDF** file using the file uploader.
    2. The content of the PDF will be extracted and displayed.
    3. Ask a query related to the PDF to get context-based answers.
    4. Download the response in **SCORM-compliant** formats.
        """)
    
st.header("📄 PDF Analysis")

# Upload PDF
pdf_file = st.file_uploader("Upload a PDF", type="pdf")
if pdf_file:
    # Extract text from the uploaded PDF
    with io.BytesIO(pdf_file.read()) as pdf_stream:
        extracted_text = extract_text_from_pdf(pdf_stream)

    # Display extracted text
    st.write("Extracted Text:")
    st.text_area("PDF Content", extracted_text, height=200)

    # Input field to ask a question
    query = st.text_input("Ask a question based on the PDF:")
    if query:
        # Generate GPT response based on the PDF content
        response = fetch_gpt_response(f"Context: {extracted_text}\nQuestion: {query}")

        # Display the generated response
        st.subheader("Response")
        st.write(response)

        # Display download options
        st.subheader("Download Options")

        # Button to download the PDF content and response as a SCORM package
        scorm_button = st.button("Generate the Response as PDF SCORM Package")
        if scorm_button:
            # Only save the response in the SCORM package
            save_as_scorm_pdf(response)  # Save only the response in PDF
            st.success("SCORM package generated. Check the 'Download SCORM Package' button.")

        if st.button("Generate the Response as Word SCORM Package"):
            # Generate the SCORM package for the Word document with only the response
            scorm_word = save_as_scorm_word(response, file_name="response.docx")

            if scorm_word:
                # Display success message
                st.success("SCORM package generated. Click the 'Download SCORM Package' button below.")

                # Display the download button
                st.download_button(
                    label="Download SCORM Word Package",
                    data=scorm_word,
                    file_name="scorm_word_package.zip",
                    mime="application/zip"
                    )
            else:
                st.error("Failed to generate SCORM Word package.")


# Horizontal line
st.markdown("---")

# Footer
st.caption("Developed by **Corbin Technology Solutions**")



    
